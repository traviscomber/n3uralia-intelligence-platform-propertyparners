#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_heading_centered(doc, text, level=1):
    """Add centered heading"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def add_table_with_borders(doc, rows, cols):
    """Add table with borders"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '12')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    
    return table

def create_contract():
    """Generate professional Word contract document"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Title
    add_heading_centered(doc, 'CONTRATO DE TRABAJO', level=1)
    add_heading_centered(doc, 'ENTRE N3URALIA Y PROPERTY PARTNERS GROUP', level=2)
    
    doc.add_paragraph()  # Spacing
    
    # Parties section
    add_heading_centered(doc, 'PARTE CONTRATANTE', level=2)
    
    p = doc.add_paragraph()
    p.add_run('EMPRESA: N3URALIA (La Empresa)\n').bold = True
    p.add_run('• Razón Social: ').bold = True
    p.add_run('n3uralia - Inteligencia Artificial y Automatización\n')
    p.add_run('• Giro: ').bold = True
    p.add_run('Desarrollo, implementación y soporte de plataformas de inteligencia artificial para gestión inmobiliaria\n')
    p.add_run('• Representante Legal: ').bold = True
    p.add_run('[A completar]\n')
    p.add_run('• Domicilio: ').bold = True
    p.add_run('[A completar]')
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    p.add_run('EMPRESA: PROPERTY PARTNERS GROUP (Cliente/Contraparte)\n').bold = True
    p.add_run('• Razón Social: ').bold = True
    p.add_run('Property Partners Group SpA\n')
    p.add_run('• Giro: ').bold = True
    p.add_run('Corretaje inmobiliario y gestión de propiedades\n')
    p.add_run('• Representante Legal: ').bold = True
    p.add_run('[A completar]\n')
    p.add_run('• Domicilio: ').bold = True
    p.add_run('[A completar]\n')
    p.add_run('• Email Principal: ').bold = True
    p.add_run('[A completar]')
    p.paragraph_format.space_after = Pt(12)
    
    # Section 1: Objeto del Contrato
    add_heading_centered(doc, '1. OBJETO DEL CONTRATO', level=2)
    
    p = doc.add_paragraph(
        'N3URALIA se compromete a proporcionar servicios de desarrollo, implementación, mantenimiento y soporte de la '
        'Plataforma de Inteligencia de Mercado Property Partners (en adelante, "la Plataforma"), incluyendo:'
    )
    p.paragraph_format.space_after = Pt(10)
    
    items = [
        ('Desarrollo de Dashboards Ejecutivos:', [
            'Dashboard CEO (métricas globales, ranking de directores, análisis de tendencias)',
            'Dashboard Director (gestión de equipos, seguimiento de agentes)',
            'Dashboard Agente (actividades diarias, pipeline personal, rankings)'
        ]),
        ('Gestión de Base de Datos:', [
            'Diseño, implementación y mantención de esquema PostgreSQL',
            'Integración con Supabase',
            'Backup y recuperación de datos'
        ]),
        ('Web Scraping de Propiedades:', [
            'Integración con múltiples portales inmobiliarios',
            'Actualización automatizada de inventario',
            'Validación de datos'
        ]),
        ('Análisis de Mercado:', [
            'Generación de reportes inteligentes con IA',
            'Análisis de precio por m² y velocidad de venta',
            'Segmentación por zonas y tipos de propiedad'
        ]),
        ('Administración de Usuarios:', [
            'Sistema de autenticación y roles',
            'Gestión de perfiles (CEO, Director, Agente)',
            'Control de acceso por rol'
        ]),
        ('Soporte Técnico:', [
            'Mantenimiento de servidor',
            'Resolución de errores y bugs',
            'Actualizaciones de seguridad',
            'Disponibilidad 24/7 para incidentes críticos'
        ])
    ]
    
    for title, subitems in items:
        p = doc.add_paragraph(title, style='List Bullet')
        p.runs[0].bold = True
        for item in subitems:
            doc.add_paragraph(item, style='List Bullet 2')
    
    doc.add_paragraph()  # Spacing
    
    # Section 2: Vigencia
    add_heading_centered(doc, '2. VIGENCIA DEL CONTRATO', level=2)
    
    p = doc.add_paragraph()
    p.add_run('• Fecha de Inicio: ').bold = True
    p.add_run('[Fecha a completar]\n')
    p.add_run('• Duración: ').bold = True
    p.add_run('[Especificar: 1 año, 3 años, indefinido, etc.]\n')
    p.add_run('• Renovación: ').bold = True
    p.add_run('Se renovará automáticamente por períodos iguales, salvo notificación contraria con 30 días de anticipación')
    p.paragraph_format.space_after = Pt(12)
    
    # Section 3: Remuneración
    add_heading_centered(doc, '3. REMUNERACIÓN Y FORMA DE PAGO', level=2)
    
    doc.add_heading('3.1 Honorarios Mensuales', level=3)
    
    table = add_table_with_borders(doc, 5, 3)
    table.rows[0].cells[0].text = 'Concepto'
    table.rows[0].cells[1].text = 'Monto'
    table.rows[0].cells[2].text = 'Moneda'
    
    table.rows[1].cells[0].text = 'Servicios Base (Desarrollo + Mantenimiento)'
    table.rows[1].cells[1].text = '[A completar]'
    table.rows[1].cells[2].text = 'UF/CLP/USD'
    
    table.rows[2].cells[0].text = 'Soporte Técnico 24/7'
    table.rows[2].cells[1].text = '[A completar]'
    table.rows[2].cells[2].text = 'UF/CLP/USD'
    
    table.rows[3].cells[0].text = 'Alojamiento Servidor (Vercel)'
    table.rows[3].cells[1].text = '[A completar]'
    table.rows[3].cells[2].text = 'UF/CLP/USD'
    
    table.rows[4].cells[0].text = 'TOTAL MENSUAL'
    table.rows[4].cells[1].text = '[A completar]'
    table.rows[4].cells[2].text = 'UF/CLP/USD'
    
    # Make header row bold
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    doc.add_paragraph()  # Spacing
    
    doc.add_heading('3.2 Condiciones de Pago', level=3)
    
    p = doc.add_paragraph()
    p.add_run('• Plazo: ').bold = True
    p.add_run('30 días desde emisión de factura\n')
    p.add_run('• Forma de pago: ').bold = True
    p.add_run('Transferencia bancaria\n')
    p.add_run('• Facturación: ').bold = True
    p.add_run('Mensual, emitida el primer día del mes\n')
    p.add_run('• Multa por atraso: ').bold = True
    p.add_run('1.5% mensual sobre saldo vencido\n')
    p.add_run('• Descuento por pago anticipado: ').bold = True
    p.add_run('[A negociar]')
    p.paragraph_format.space_after = Pt(12)
    
    # Section 4: Obligaciones N3URALIA
    add_heading_centered(doc, '4. OBLIGACIONES DE N3URALIA', level=2)
    
    obligations = [
        ('Proveer Servicios de Calidad:', [
            'Código limpio, documentado y testeado',
            'Cumplimiento con estándares de seguridad (OWASP, SSL/TLS)',
            'Disponibilidad mínima del 99.5% de la Plataforma'
        ]),
        ('Mantener la Plataforma:', [
            'Actualizaciones de seguridad dentro de 48 horas de identificación',
            'Parches mensuales de mantenimiento',
            'Monitoreo proactivo de performance'
        ]),
        ('Soporte Técnico:', [
            'Respuesta a incidentes dentro de 2 horas (críticos), 24 horas (moderados)',
            'Documentación técnica completa',
            'Capacitación inicial del equipo de PP Group'
        ]),
        ('Confidencialidad:', [
            'No compartir datos de PP Group con terceros',
            'Usar datos solo para propósitos contractuales',
            'Cumplir con LGPD/GDPR en caso de datos personales'
        ]),
        ('Entrega de Acceso:', [
            'Proporcionar credenciales seguras',
            'Mantener documentación de usuarios activos',
            'Restaurar acceso cuando sea necesario'
        ])
    ]
    
    for title, items in obligations:
        p = doc.add_paragraph(title, style='List Bullet')
        p.runs[0].bold = True
        for item in items:
            doc.add_paragraph(item, style='List Bullet 2')
    
    doc.add_paragraph()  # Spacing
    
    # Section 5: Obligaciones PP Group
    add_heading_centered(doc, '5. OBLIGACIONES DE PROPERTY PARTNERS GROUP', level=2)
    
    pp_obligations = [
        ('Pagar Puntualmente:', [
            'Enviar pagos según cronograma establecido',
            'Notificar de cualquier inconveniente de pago con anticipación'
        ]),
        ('Comunicación Efectiva:', [
            'Designar persona de contacto para coordinación',
            'Reportar problemas o errores rápidamente',
            'Proporcionar feedback sobre nuevas features'
        ]),
        ('Datos Precisos:', [
            'Mantener información de usuarios actualizada',
            'Proporcionar datos de prueba representativos',
            'Avisar sobre cambios en estructura organizacional'
        ]),
        ('Respeto de Términos:', [
            'No intentar acceso no autorizado',
            'No compartir credenciales con terceros',
            'No reversa-ingenierizar o copiar código'
        ])
    ]
    
    for title, items in pp_obligations:
        p = doc.add_paragraph(title, style='List Bullet')
        p.runs[0].bold = True
        for item in items:
            doc.add_paragraph(item, style='List Bullet 2')
    
    # Add page break before more sections
    doc.add_page_break()
    
    # Section 6: Propiedad Intelectual
    add_heading_centered(doc, '6. PROPIEDAD INTELECTUAL', level=2)
    
    items = [
        ('Código Fuente:', 'PP Group obtiene licencia de uso perpetua del código (sin derecho de transferencia). N3URALIA retiene propiedad del código fuente. En caso de terminación, PP Group mantiene acceso por 30 días para migración.'),
        ('Datos:', 'PP Group es propietario de todos los datos generados en su uso. N3URALIA puede usar datos agregados/anonimizados para mejora de productos.'),
        ('Derivados:', 'Cualquier módulo personalizado pertenece a PP Group. N3URALIA puede usar técnicas/conceptos para futuros clientes.')
    ]
    
    for title, desc in items:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        p.add_run(' ' + desc)
        p.paragraph_format.space_after = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # Section 7: Confidencialidad
    add_heading_centered(doc, '7. CONFIDENCIALIDAD Y SEGURIDAD', level=2)
    
    doc.add_heading('7.1 Datos Protegidos', level=3)
    p = doc.add_paragraph('Se consideran confidenciales:')
    for item in ['Información de propiedades y precios', 'Datos de clientes y contactos', 'Estrategias comerciales', 'Información financiera']:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('7.2 Medidas de Seguridad', level=3)
    p = doc.add_paragraph('N3URALIA implementará:')
    for item in ['Encriptación SSL/TLS en tránsito', 'Encriptación a nivel base de datos', 'Autenticación multi-factor (MFA)', 'Auditoría de accesos', 'Backups diarios con retención de 30 días']:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('7.3 Período de Confidencialidad', level=3)
    p = doc.add_paragraph('Tres (3) años después de terminación del contrato.')
    p.paragraph_format.space_after = Pt(12)
    
    # Section 8: Límites de Responsabilidad
    add_heading_centered(doc, '8. LÍMITES DE RESPONSABILIDAD', level=2)
    
    doc.add_heading('8.1 Limitación de Daños', level=3)
    p = doc.add_paragraph(
        'En caso de incumplimiento, la responsabilidad de N3URALIA se limita al 50% del monto pagado en los últimos 12 meses, '
        'excepto en casos de violación grave de confidencialidad, negligencia grave que cause daño directo, o pérdida de datos por acción deliberada.'
    )
    p.paragraph_format.space_after = Pt(12)
    
    doc.add_heading('8.2 Exclusiones', level=3)
    p = doc.add_paragraph('N3URALIA no es responsable por:')
    for item in ['Datos perdidos por cliente (no uso de backups)', 'Indisponibilidad por falta de pago', 'Incompatibilidad con software de terceros', 'Uso incorrecto de la Plataforma', 'Ataques cibernéticos exitosos (se excluyen negligencias)']:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # Section 9: Terminación
    add_heading_centered(doc, '9. TERMINACIÓN', level=2)
    
    doc.add_heading('9.1 Por Causa', level=3)
    p = doc.add_paragraph('Cualquier parte puede terminar sin costo si:')
    for item in ['La otra parte incumple obligaciones graves por más de 15 días', 'Se declara quiebra o insolvencia', 'Se comete fraude']:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('9.2 Por Conveniencia', level=3)
    items = ['PP Group: Puede terminar con 90 días de notificación, sin costo', 'N3URALIA: Puede terminar con 60 días de notificación, sin costo']
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('9.3 Al Vencimiento', level=3)
    p = doc.add_paragraph('Si no se renueva, se otorgan 30 días de acceso de lectura para exportar datos.')
    p.paragraph_format.space_after = Pt(12)
    
    # Section 10: Transición
    add_heading_centered(doc, '10. TRANSICIÓN Y MIGRACIÓN', level=2)
    p = doc.add_paragraph('En caso de terminación:')
    for item in ['N3URALIA exportará toda la base de datos en formato estándar (CSV, JSON)', 'Proporciona 30 días de acceso de lectura', 'N3URALIA borrará todos los datos después de 60 días', 'Documentación técnica completa será entregada']:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_paragraph()
    
    # Section 11-14: Other sections (abbreviated for space)
    add_heading_centered(doc, '11. ESCALABILIDAD Y CRECIMIENTO', level=2)
    doc.add_paragraph('La Plataforma será diseñada para soportar hasta 500 usuarios concurrentes inicialmente, crecimiento a 2,000+ usuarios sin rediseño, escalamiento automático de infraestructura, y performance con carga de páginas menor a 2 segundos.')
    
    add_heading_centered(doc, '12. CUMPLIMIENTO NORMATIVO', level=2)
    p = doc.add_paragraph('Las partes acuerdan cumplir con:')
    for item in ['Leyes de protección de datos vigentes en Chile', 'Normas de seguridad ISO 27001', 'Estándares de accesibilidad WCAG 2.1', 'Regulaciones de comercio electrónico aplicables']:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_centered(doc, '13. RESOLUCIÓN DE CONTROVERSIAS', level=2)
    items = ['Negociación: Las partes buscarán resolver disputas directamente', 'Mediación: Si no hay acuerdo, se recurrirá a mediador', 'Arbitraje: Según reglas de [Cámara Comercio/Centro Arbitraje]', 'Jurisdicción: Cortes de [Ciudad], Chile']
    for i, item in enumerate(items, 1):
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # Section 14: Final
    add_heading_centered(doc, '14. DISPOSICIONES FINALES', level=2)
    
    doc.add_heading('14.1 Enmiendas', level=3)
    p = doc.add_paragraph('Cualquier modificación debe ser por escrito y firmada por ambas partes.')
    p.paragraph_format.space_after = Pt(10)
    
    doc.add_heading('14.2 Fuerza Mayor', level=3)
    p = doc.add_paragraph(
        'Ninguna parte es responsable por incumplimientos por casos fortuitos (pandemias, desastres naturales, etc.) '
        'siempre que notifique dentro de 48 horas.'
    )
    p.paragraph_format.space_after = Pt(10)
    
    doc.add_heading('14.3 Integralidad', level=3)
    p = doc.add_paragraph('Este contrato constituye el acuerdo completo entre las partes, reemplazando cualquier acuerdo previo.')
    p.paragraph_format.space_after = Pt(10)
    
    doc.add_heading('14.4 Divisibilidad', level=3)
    p = doc.add_paragraph('Si alguna cláusula es inválida, las demás permanecen en vigor.')
    p.paragraph_format.space_after = Pt(12)
    
    # Signatures
    doc.add_paragraph()
    add_heading_centered(doc, 'FIRMAS', level=2)
    
    doc.add_paragraph()
    p = doc.add_paragraph(f'En Santiago de Chile, a los _______ de __________________ de 202__.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    add_heading_centered(doc, 'Por N3URALIA:', level=3)
    
    doc.add_paragraph()
    doc.add_paragraph('_____________________________')
    doc.add_paragraph('Nombre:')
    doc.add_paragraph('Cargo:')
    doc.add_paragraph('Cédula de Identidad:')
    doc.add_paragraph('Fecha:')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_heading_centered(doc, 'Por PROPERTY PARTNERS GROUP:', level=3)
    
    doc.add_paragraph()
    doc.add_paragraph('_____________________________')
    doc.add_paragraph('Nombre:')
    doc.add_paragraph('Cargo:')
    doc.add_paragraph('Cédula de Identidad:')
    doc.add_paragraph('Fecha:')
    
    # Footer with legal notes
    doc.add_page_break()
    
    p = doc.add_paragraph()
    p.add_run('NOTAS LEGALES\n').bold = True
    p.add_run('⚠️ IMPORTANTE:\n').bold = True
    p.add_run(
        'Este documento es un TEMPLATE y requiere ser revisado y adaptado por abogados especializados en derecho comercial '
        'antes de su ejecución.\n\n'
        'Debe ser personalizado con:\n'
        '• Datos específicos de ambas empresas\n'
        '• Montos y condiciones comerciales finales\n'
        '• Jurisdicción y leyes aplicables específicas\n'
        '• Información de contactos y representantes legales\n\n'
        'No constituye asesoramiento legal.'
    )
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Versión: ').bold = True
    p.add_run('1.0\n')
    p.add_run('Última Actualización: ').bold = True
    p.add_run('Julio 2026\n')
    p.add_run('Documento Preparado Por: ').bold = True
    p.add_run('n3uralia Legal Template')
    
    # Save document
    output_path = '/vercel/share/v0-project/CONTRATO_TRABAJO_n3uralia_ppgroup.docx'
    doc.save(output_path)
    
    print(f'[v0] ✅ Contrato generado exitosamente')
    print(f'[v0] Ubicación: {output_path}')
    print(f'[v0] Tamaño: {os.path.getsize(output_path) / 1024:.1f} KB')

if __name__ == '__main__':
    import os
    create_contract()
